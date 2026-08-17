# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
WORK = ROOT / "tmp" / "report_work"
OUTPUT = ROOT / "output"
TEMPLATE = WORK / "template.docx"
OUT_DOCX = OUTPUT / "한이음_드림업_중간_개발보고서_초안.docx"
ARCH_IMG = WORK / "architecture_flow.png"


def ensure_output() -> None:
    OUTPUT.mkdir(exist_ok=True)


def set_run_font(run, size: float = 10, bold: bool = False) -> None:
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Malgun Gothic")
    rfonts.set(qn("w:hAnsi"), "Malgun Gothic")
    rfonts.set(qn("w:eastAsia"), "맑은 고딕")


def clear_paragraph(paragraph) -> None:
    paragraph._p.clear_content()


def write_paragraph(paragraph, text: str, size: float = 10, bold: bool = False) -> None:
    clear_paragraph(paragraph)
    lines = text.split("\n")
    for idx, line in enumerate(lines):
        if idx:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=size, bold=bold)


def add_after(paragraph, text: str, style: str | None = None, size: float = 10, bold: bool = False):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph.__class__(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    write_paragraph(new_para, text, size=size, bold=bold)
    return new_para


def set_cell_text(cell, text: str, size: float = 9.5, bold: bool = False, align=None) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    write_paragraph(p, text, size=size, bold=bold)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def delete_table(table) -> None:
    element = table._element
    element.getparent().remove(element)


def delete_row(table, row_index: int) -> None:
    tr = table.rows[row_index]._tr
    tr.getparent().remove(tr)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def table_text(table) -> str:
    texts: list[str] = []
    seen = set()
    for row in table.rows:
        for cell in row.cells:
            if id(cell._tc) in seen:
                continue
            seen.add(id(cell._tc))
            txt = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            if txt:
                texts.append(txt)
    return "\n".join(texts)


def insert_picture_cell(cell, image_path: Path, caption: str, width: float = 1.45) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_path.exists():
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(width))
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        write_paragraph(p2, caption, size=8.2)
    else:
        write_paragraph(p, caption, size=8.5)


def make_architecture_image() -> None:
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return

    width, height = 1600, 380
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 29)
        small = ImageFont.truetype("arial.ttf", 23)
    except Exception:
        font = ImageFont.load_default()
        small = ImageFont.load_default()

    boxes = [
        ("Camera\nReed Switch", 35, 95, 285, 255, "#E8F2FF"),
        ("Raspberry Pi\nBridge", 355, 95, 605, 255, "#EAF7EE"),
        ("YOLO\nClassifier", 675, 95, 925, 255, "#FFF4D9"),
        ("Flask API\nMySQL", 995, 95, 1245, 255, "#F3ECFF"),
        ("Flutter App\nInventory\nRecipe", 1315, 95, 1570, 255, "#FFECEC"),
    ]
    for label, x1, y1, x2, y2, color in boxes:
        draw.rounded_rectangle([x1, y1, x2, y2], radius=24, fill=color, outline="#3D4A5C", width=3)
        lines = label.split("\n")
        total_h = len(lines) * 34
        for idx, line in enumerate(lines):
            bbox = draw.textbbox((0, 0), line, font=font)
            tx = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
            ty = y1 + (y2 - y1 - total_h) / 2 + idx * 37
            draw.text((tx, ty), line, fill="#111827", font=font)

    for i in range(len(boxes) - 1):
        _, _x1, _y1, x2, _y2, _ = boxes[i]
        _, nx1, _ny1, _nx2, _ny2, _ = boxes[i + 1]
        y = 175
        draw.line([x2 + 18, y, nx1 - 18, y], fill="#374151", width=5)
        draw.polygon([(nx1 - 18, y), (nx1 - 38, y - 12), (nx1 - 38, y + 12)], fill="#374151")

    draw.text((45, 315), "Event-driven capture -> AI recognition -> inventory DB -> user-facing management", fill="#4B5563", font=small)
    ARCH_IMG.parent.mkdir(exist_ok=True)
    img.save(ARCH_IMG)


def fill_summary_tables(doc: Document) -> None:
    tables = doc.tables

    set_cell_text(tables[2].cell(0, 1), "독거노인의 식생활 관리를 위한 AI 기반 스마트 냉장고", size=11, bold=True)
    set_cell_text(tables[3].cell(0, 0), "요 약 본", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    info = tables[4]
    set_cell_text(info.cell(1, 1), "독거노인의 식생활 관리를 위한 AI 기반 스마트 냉장고", size=9.5, bold=True)
    set_cell_text(info.cell(2, 1), "☑ 생활", size=9)
    set_cell_text(info.cell(2, 2), "□ 업무", size=9)
    set_cell_text(info.cell(2, 3), "□ 공공/교통", size=9)
    set_cell_text(info.cell(2, 4), "□ 금융/핀테크", size=9)
    set_cell_text(info.cell(3, 1), "□ 의료", size=9)
    set_cell_text(info.cell(3, 2), "□ 교육", size=9)
    set_cell_text(info.cell(3, 3), "□ 유통/쇼핑", size=9)
    set_cell_text(info.cell(3, 4), "□ 엔터테인먼트", size=9)
    set_cell_text(info.cell(4, 1), "☑ 소프트웨어", size=9)
    set_cell_text(info.cell(4, 2), "☑ 인공지능", size=9)
    set_cell_text(info.cell(4, 3), "☑ 스마트 디바이스", size=9)
    set_cell_text(info.cell(4, 4), "□ 방송, 콘텐츠", size=9)
    set_cell_text(info.cell(5, 1), "□ 디지털융합", size=9)
    set_cell_text(info.cell(5, 2), "□ 차세대통신", size=9)
    set_cell_text(info.cell(5, 3), "□ 사이버보안", size=9)
    set_cell_text(
        info.cell(6, 1),
        "☑ 논문게재 및 포스터 발표  □ 앱등록  □ 프로그램등록  □ 특허  □ 기술이전\n"
        "□ 실용화  ☑ 공모전(한이음 드림업 공모전)  □ 기타( )",
        size=8.7,
    )
    set_cell_text(
        info.cell(7, 1),
        "냉장고 내부 카메라로 식재료 이미지를 촬영하고 AI 모델이 식재료를 자동 인식한 뒤, "
        "모바일 앱에서 보유 식재료와 추천 레시피를 확인할 수 있도록 하는 스마트 냉장고 서비스이다.",
        size=9,
    )
    set_cell_text(
        info.cell(8, 1),
        "독거노인은 식재료 보유 현황을 놓치기 쉽고, 이로 인해 식사 준비 부담과 음식물 폐기가 발생할 수 있다. "
        "자동 인식과 재고 관리를 통해 식생활 관리 부담을 낮추는 것이 필요하다.",
        size=9,
    )
    set_cell_text(
        info.cell(9, 1),
        "문 열림 감지 기반 촬영, YOLO/분류 모델 기반 식재료 인식, 재고 자동 등록/소비 처리, "
        "보유 식재료 기반 레시피 추천을 하나의 앱 서비스로 연결한다.",
        size=9,
    )
    set_cell_text(
        info.cell(10, 1),
        "회원/냉장고 관리, 식재료 이미지 인식 및 업로드, 재고 목록 조회/수정/삭제, "
        "보유 식재료 기준 레시피 추천, Raspberry Pi 카메라 브리지 및 리드스위치 연동.",
        size=9,
    )
    set_cell_text(
        info.cell(11, 1),
        "독거노인 식생활 지원, 음식물 낭비 감소, 스마트홈/돌봄/생활관리 서비스 확장, "
        "AI 기반 주방 재고 관리 연구 및 공모전 결과물로 활용 가능하다.",
        size=9,
    )

    set_cell_text(tables[5].cell(0, 0), "본    문", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def fill_body_paragraphs(doc: Document) -> None:
    p = doc.paragraphs
    write_paragraph(p[15], "I. 프로젝트 개요", size=13, bold=True)
    write_paragraph(p[18], "   1) 프로젝트 정의", size=10, bold=True)
    write_paragraph(
        p[19],
        "     - 본 프로젝트는 냉장고 내부 카메라와 AI 이미지 인식 모델을 활용하여 식재료를 자동으로 식별하고, "
        "인식 결과를 데이터베이스와 모바일 앱으로 연결하는 스마트 냉장고 관리 시스템이다.",
        size=10,
    )
    last = p[19]
    for text in [
        "     - 사용자는 앱에서 냉장고별 식재료 재고를 확인하고, 보유 식재료를 기반으로 추천 레시피와 부족한 재료를 확인할 수 있다.",
        "     - Raspberry Pi는 카메라 촬영, 후보 영역 탐지, 분류 결과 업로드를 담당하고 Flask 백엔드는 인증, 냉장고, 재고, 레시피 API를 제공한다.",
    ]:
        last = add_after(last, text, style="소제목")

    write_paragraph(p[23], "   1) 개발 동기", size=10, bold=True)
    write_paragraph(
        p[24],
        "     - 고령화로 독거노인 가구가 증가하면서 식재료 관리, 식사 준비, 음식물 폐기 문제가 생활 관리의 부담으로 이어지고 있다.",
        size=10,
    )
    last = p[24]
    for text in [
        "     - 냉장고 안의 식재료를 직접 기억하거나 수기로 관리해야 하는 방식은 고령층에게 번거롭고 누락 가능성이 크다.",
        "     - 자동 촬영 및 인식 기반 재고 관리는 복잡한 입력 없이 식재료 현황을 파악하게 하므로 스마트홈 기반 돌봄 서비스로 확장할 수 있다.",
    ]:
        last = add_after(last, text, style="소제목")

    write_paragraph(p[28], "   1) 기능적·기술적 차별성", size=10, bold=True)
    write_paragraph(
        p[29],
        "     - 기존의 수동 냉장고 관리 앱과 달리 카메라 촬영, AI 인식, 재고 저장, 레시피 추천을 하나의 흐름으로 연결한다.",
        size=10,
    )
    last = p[29]
    for text in [
        "     - 고정된 중앙 영역만 분류하지 않고 YOLO 탐지 박스, 윤곽선 기반 후보, 중앙 후보를 순차 적용하여 실제 냉장고 환경의 배치 변화에 대응한다.",
        "     - 리드스위치 기반 문 열림/닫힘 이벤트를 활용하여 식재료 추가와 소비 처리를 자동화할 수 있도록 설계했다.",
        "     - 앱, 백엔드, 데이터베이스, Raspberry Pi 브리지 문서를 분리해 재현성과 유지보수성을 높였다.",
    ]:
        last = add_after(last, text, style="소제목")

    write_paragraph(p[31], "II. 프로젝트 내용", size=13, bold=True)
    write_paragraph(p[34], "   1) 전체 시스템 구성", size=10, bold=True)
    write_paragraph(
        p[35],
        "     - 냉장고 내부 카메라와 리드스위치가 사용자 동작을 감지하고, Raspberry Pi가 이미지를 수집하여 AI 모델과 백엔드로 전달한다.",
        size=10,
    )

    write_paragraph(p[48], "     - S/W는 Flutter 앱, Flask REST API, MySQL 데이터베이스, YOLO 기반 인식 파이프라인으로 구성된다.", size=10)
    write_paragraph(p[53], "     - H/W는 Raspberry Pi 5, 카메라 모듈, 리드스위치, 냉장고 내부 촬영 환경을 중심으로 구성된다.", size=10)

    write_paragraph(p[55], "3. 주요 적용 기술", size=11, bold=True)
    write_paragraph(
        p[56],
        "     - AI 이미지 인식: Ultralytics YOLOv8n 탐지 모델과 식재료 분류 모델(best.pt)을 조합하여 후보 영역과 식재료명을 산출한다.",
        size=10,
    )
    last = p[56]
    for text in [
        "     - 영상 처리: OpenCV 기반 윤곽선 후보 탐지, crop padding, confidence threshold, stable frame 판정을 적용하여 오인식을 줄인다.",
        "     - 백엔드/DB: Flask REST API와 MySQL을 사용해 사용자, 냉장고, 식재료, 재고, 레시피, 업로드 이미지를 관리한다.",
        "     - 모바일 앱: Flutter와 HTTP 통신으로 로그인, 냉장고 선택, 재고 CRUD, 레시피 추천 화면을 구현한다.",
        "     - IoT 연동: Picamera2/OpenCV 카메라, gpiozero 리드스위치, MJPEG preview stream으로 Raspberry Pi 현장 테스트를 지원한다.",
    ]:
        last = add_after(last, text, style="소제목")

    write_paragraph(p[60], " 5. 기타 사항", size=11, bold=True)
    write_paragraph(
        p[61],
        "     - 본 프로젝트는 단순한 앱 화면 구현에 머무르지 않고, 실제 카메라 입력부터 AI 분류, API 저장, 앱 조회까지 이어지는 end-to-end 흐름을 목표로 한다.",
        size=10,
    )
    last = p[61]
    for text in [
        "     - 고령층 사용자를 고려하여 복잡한 수기 입력을 줄이고 자동 인식 중심의 재고 관리 경험을 제공한다.",
        "     - 향후 실물 장착 안정화, 데이터셋 확장, 인식 정확도 검증을 통해 복지·스마트홈 서비스로 확장할 수 있다.",
    ]:
        last = add_after(last, text, style="소제목")

    write_paragraph(p[63], "III. 프로젝트 수행 내용", size=13, bold=True)
    write_paragraph(p[65], "프로젝트 수행일정", size=11, bold=True)
    write_paragraph(p[70], "  1) 프로젝트 관리 측면", size=10, bold=True)
    write_paragraph(
        p[71],
        "     - 문제점: 앱, 백엔드, AI 모델, Raspberry Pi가 서로 다른 환경에서 개발되어 네트워크 주소와 실행 절차가 혼동될 수 있었다.",
        size=10,
    )
    last = p[71]
    for text in [
        "     - 해결방안: Git/GitHub 기반 형상관리와 HANDOFF.md, RASPBERRY_PI.md 문서를 작성하여 백엔드 실행, Pi 연결, 테스트 명령을 표준화했다.",
        "     - 일정 관리는 주 1회 진행 공유를 기준으로 하며, H/W 연동 리스크가 큰 부분은 dry-run, preview, once 실행 옵션으로 단계별 검증이 가능하도록 했다.",
    ]:
        last = add_after(last, text, style="소제목")

    write_paragraph(p[75], "  2) 프로젝트 개발 측면", size=10, bold=True)
    write_paragraph(
        p[76],
        "     - 문제점: 냉장고 내부 식재료 위치가 고정되어 있지 않아 중앙 crop만으로는 실제 사용 환경에서 인식률이 떨어질 수 있었다.",
        size=10,
    )
    last = p[76]
    for text in [
        "     - 해결방안: YOLO 탐지 박스, 윤곽선 후보, 중앙 fallback을 조합하고 crop padding을 적용해 식재료 후보를 동적으로 추출하도록 개선했다.",
        "     - 문제점: 문 열림 시마다 중복 업로드가 발생할 수 있어 재고 데이터가 부정확해질 가능성이 있었다.",
        "     - 해결방안: stable frame 판정, upload cooldown, add-on-open/consume-on-close workflow를 적용해 반복 업로드와 소비 처리를 제어했다.",
        "     - 남은 과제: 실제 냉장고 장착 사진 확보, 다양한 식재료 데이터셋 추가 학습, 모바일 UI 최종 정리, 정확도/응답시간 테스트가 필요하다.",
    ]:
        last = add_after(last, text, style="소제목")

    write_paragraph(p[80], "  3. 프로젝트를 통해 배우거나 느낀 점", size=11, bold=True)
    write_paragraph(
        p[81],
        "     - AI 모델 단독 구현보다 실제 서비스에서는 촬영 조건, 후보 영역 추출, API 저장, 앱 표시까지 연결되는 전체 파이프라인 안정성이 중요함을 배웠다.",
        size=10,
    )
    last = p[81]
    for text in [
        "     - Raspberry Pi와 서버가 같은 네트워크에서 통신해야 하는 점, 카메라 FPS와 AI 추론 주기의 균형 등 H/W 연동에서 고려해야 할 요소를 경험했다.",
        "     - 사용자 관점에서는 기능 수보다 자동화 흐름의 신뢰성, 결과 확인의 쉬움, 오류 발생 시 수정 가능한 UI가 중요하다는 점을 확인했다.",
    ]:
        last = add_after(last, text, style="소제목")

    write_paragraph(p[98], "Ⅳ. 기대효과 및 활용분야", size=13, bold=True)
    write_paragraph(p[100], "     - 독거노인이 보유 식재료를 쉽게 확인하고, 식재료 기반 추천 레시피를 통해 식사 준비 부담을 줄일 수 있다.", size=10)
    last = p[100]
    for text in [
        "     - 식재료 활용도를 높여 음식물 폐기를 줄이고, 장보기·식사 계획을 더 체계적으로 관리할 수 있다.",
        "     - 수동 입력을 최소화한 자동 인식 방식이므로 고령층 친화형 스마트 주방 서비스로 발전 가능하다.",
        "     - 서비스 데이터를 축적하면 식재료 소비 패턴 분석, 맞춤형 식단 추천, 돌봄 알림 서비스로 확장할 수 있다.",
    ]:
        last = add_after(last, text, style="소제목")

    write_paragraph(p[107], "     - 독거노인 및 1인 가구의 냉장고 재고 관리 서비스", size=10)
    last = p[107]
    for text in [
        "     - 지역 복지기관, 돌봄 서비스, 스마트홈 플랫폼과 연계한 생활 관리 솔루션",
        "     - 식재료 인식 데이터셋 구축 및 AI 이미지 분류 연구",
        "     - 냉장고, 주방가전, 모바일 앱을 연결하는 스마트 주방 서비스",
    ]:
        last = add_after(last, text, style="소제목")


def fill_project_tables(doc: Document) -> None:
    tables = doc.tables

    # Architecture diagram box.
    make_architecture_image()
    arch_table = tables[11]
    cell = arch_table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_paragraph(p, "전체 서비스 흐름도", size=10, bold=True)
    if ARCH_IMG.exists():
        pic_p = cell.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_p.add_run().add_picture(str(ARCH_IMG), width=Inches(6.1))
    desc = cell.add_paragraph()
    write_paragraph(
        desc,
        "냉장고 내부 촬영/문 열림 감지 → Raspberry Pi 브리지 → YOLO 탐지 및 식재료 분류 → Flask API/MySQL 저장 → Flutter 앱 조회·추천",
        size=9,
    )

    functions = tables[12]
    function_rows = [
        ("S/W", "AI 식재료 인식", "YOLO 탐지 후보와 식재료 분류 모델을 사용해 이미지 속 식재료를 판별하고 confidence와 crop 이미지를 저장한다. 8/31까지 모델 검증 및 클래스 확장 예정.", "75"),
        ("S/W", "재고/레시피 앱", "Flutter 앱에서 로그인, 냉장고 선택, 재고 조회/추가/수정/삭제, 보유 재료 기반 레시피 추천과 부족 재료 확인을 제공한다. 8/20까지 UI 보완 예정.", "70"),
        ("H/W", "카메라 브리지", "Raspberry Pi가 640x480, 30FPS 카메라 프레임을 읽고 후보 영역을 분류한 뒤 백엔드 /upload로 전송한다. 9/5까지 실물 장착 테스트 예정.", "65"),
        ("H/W", "문 열림/소비 감지", "리드스위치 GPIO17 기반으로 문 열림 시 추가, 문 닫힘 시 소비 처리 workflow를 구현 중이다. 9/10까지 배선/오동작 테스트 예정.", "50"),
    ]
    for row_idx, row_data in enumerate(function_rows, start=1):
        for col_idx, value in enumerate(row_data):
            set_cell_text(functions.cell(row_idx, col_idx), value, size=8.4 if col_idx == 2 else 9)
    delete_row(functions, 5)

    sw_table = tables[13]
    sw_rows = [
        ("식재료 인식/업로드", "카메라 이미지에서 식재료 후보를 탐지·분류하고 원본/크롭 이미지, 라벨, confidence를 Flask 백엔드에 업로드한다."),
        ("재고 관리", "사용자는 앱에서 냉장고별 재고 목록을 보고 수량, 단위, 상태, 메모를 직접 수정하거나 삭제할 수 있다."),
        ("레시피 추천", "현재 냉장고 재고와 레시피 필요 재료를 비교하여 만들 수 있는 레시피와 부족한 재료를 함께 표시한다."),
    ]
    images = [
        ROOT / "fridge_app" / "assets" / "images" / "sample_fruit_01.jpg",
        ROOT / "backend" / "egg_dataset_preview.jpg",
        ROOT / "fridge_app" / "assets" / "images" / "sample_fruit_02.jpg",
    ]
    captions = ["테스트 입력 이미지", "AI 학습/검증 이미지", "앱 표시용 샘플 이미지"]
    for idx, (function, desc_text) in enumerate(sw_rows, start=1):
        set_cell_text(sw_table.cell(idx, 0), function, size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(sw_table.cell(idx, 1), desc_text, size=8.2)
        insert_picture_cell(sw_table.cell(idx, 2), images[idx - 1], captions[idx - 1], width=1.1)
    delete_row(sw_table, 4)

    hw_table = tables[15]
    hw_rows = [
        ("Raspberry Pi 5 + 카메라", "냉장고 내부 식재료 촬영, AI 인식 파이프라인 실행, 백엔드 업로드를 담당한다."),
        ("리드스위치", "냉장고 문 열림/닫힘을 감지하여 식재료 추가·소비 이벤트를 자동 트리거한다."),
    ]
    for idx, (part, desc_text) in enumerate(hw_rows, start=1):
        set_cell_text(hw_table.cell(idx, 0), part, size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(hw_table.cell(idx, 1), desc_text, size=8.2)
        set_cell_text(hw_table.cell(idx, 2), "실물 장착 사진\n첨부 예정", size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    tech_box = tables[18].cell(0, 0)
    set_cell_text(
        tech_box,
        "적용 시나리오: 사용자가 냉장고 문을 열면 리드스위치 이벤트가 발생하고, Raspberry Pi가 카메라 프레임을 수집한다. "
        "YOLO 탐지 박스와 윤곽선 후보를 기반으로 식재료 crop을 만든 뒤 분류 모델이 라벨과 신뢰도를 산출한다. "
        "백엔드는 /upload 또는 /consume 요청을 받아 MySQL의 fridge_items 테이블을 갱신하고, Flutter 앱은 재고 및 레시피 API를 호출해 사용자에게 결과를 제공한다.",
        size=9.2,
    )

    env = tables[19]
    env_values = {
        1: "Windows 개발 PC, Raspberry Pi OS 64-bit, Android/Web 테스트 환경",
        2: "VS Code, Codex, Flutter 개발 환경, Python 가상환경",
        3: "Flutter, Flask, MySQL, Ultralytics YOLO, OpenCV, Git/GitHub",
        4: "Python, Dart, SQL",
        5: "REST API, CORS, 모델 가중치(yolov8n.pt, best.pt), 이미지 업로드 저장소",
        6: "Raspberry Pi 5, 카메라 모듈, 냉장고, Android 스마트폰",
        7: "리드스위치(GPIO17), 카메라",
        8: "HTTP REST 통신, LAN/Tailscale 기반 원격 접속 가능 구조, MJPEG preview stream",
        9: "Python",
        10: "640x480 30FPS 기본 촬영, warm camera/on-demand camera 모드 지원",
        11: "Git/GitHub, HANDOFF.md, RASPBERRY_PI.md",
        12: "카카오톡, 디스코드, 대면 미팅, Notion 활용 예정",
        13: "주 1회 진행 상황 공유, 단계별 dry-run/preview 테스트",
    }
    for row_idx, value in env_values.items():
        set_cell_text(env.cell(row_idx, 2), value, size=8.4)
    delete_row(env, 14)

    value_box = tables[20].cell(0, 0)
    set_cell_text(
        value_box,
        "프로젝트 가치: 고령층이 직접 입력하지 않아도 식재료 보유 현황을 파악할 수 있도록 하여 사용성을 높이고, "
        "식재료 인식 데이터와 레시피 추천을 연결해 생활 관리·음식물 폐기 감소·돌봄 서비스 연계 가능성을 확보한다. "
        "또한 탐지-분류-저장-앱 표시까지 연결된 구조라 향후 데이터셋 확장, 정확도 평가, 서비스 고도화가 용이하다.",
        size=9.2,
    )

    schedule = tables[22]
    set_cell_text(schedule.cell(1, 0), "작성 기준", size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(schedule.cell(1, 1), "2026. 7. 7. 기준", size=8.5)
    for c in range(2, 11):
        set_cell_text(schedule.cell(1, c), "", size=8)

    plans = {
        4: ("계획", "주제 선정 및 요구사항 분석", [2, 3]),
        5: ("분석", "AI/IoT 기술 조사 및 데이터셋 검토", [2, 3, 4]),
        6: ("설계", "DB/API 구조 설계", [3, 4]),
        7: ("설계", "앱 UI 및 H/W 연동 구조 설계", [3, 4, 5]),
        8: ("개발", "Flask/MySQL 백엔드 및 재고 API 구현", [4, 5]),
        9: ("개발", "Flutter 앱 화면 및 API 연동", [4, 5, 6]),
        10: ("개발", "AI 모델·Raspberry Pi 카메라 브리지 연동", [4, 5, 6, 7]),
        11: ("테스트", "인식 정확도, 앱 기능, H/W 장착 테스트", [6, 7, 8]),
        12: ("종료\n(성과등록)", "최종 결과물 정리 및 공모전/논문 자료 준비", [8]),
    }
    for row_idx, (category, task, month_cells) in plans.items():
        set_cell_text(schedule.cell(row_idx, 0), category, size=8.3, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(schedule.cell(row_idx, 1), task, size=8.2)
        for c in range(2, 11):
            mark = "완료" if c in month_cells and c <= 4 else ("진행" if c in month_cells and c == 5 else ("예정" if c in month_cells else ""))
            set_cell_text(schedule.cell(row_idx, c), mark, size=7.6, align=WD_ALIGN_PARAGRAPH.CENTER)
            if mark == "진행":
                shade_cell(schedule.cell(row_idx, c), "FFF2CC")
            elif mark == "완료":
                shade_cell(schedule.cell(row_idx, c), "E2F0D9")
            elif mark == "예정":
                shade_cell(schedule.cell(row_idx, c), "EAF2F8")


def clean_template_guidance(doc: Document) -> None:
    for table in reversed(doc.tables):
        txt = table_text(table).strip()
        if len(table.rows) == 1 and txt.startswith("#"):
            delete_table(table)

    body_started = False
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip().startswith("I. 프로젝트 개요"):
            body_started = True
        if body_started and not paragraph.text.strip():
            delete_paragraph(paragraph)


def main() -> None:
    ensure_output()
    doc = Document(str(TEMPLATE))
    fill_summary_tables(doc)
    fill_body_paragraphs(doc)
    fill_project_tables(doc)
    clean_template_guidance(doc)
    doc.save(str(OUT_DOCX))
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
