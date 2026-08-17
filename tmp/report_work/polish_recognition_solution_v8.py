# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
INPUT = ROOT / "output" / "hanium_dreamup_mid_report_senior_style_v7.docx"
OUTPUT = ROOT / "output" / "hanium_dreamup_mid_report_senior_style_v8.docx"

OLD_TEXT = "- 향후 실제 냉장고 장착 테스트와 다양한 식재료 데이터셋 추가 학습을 통해 인식 정확도를 보완할 예정이다."
NEW_TEXT = "- 향후 실제 냉장고 장착 테스트를 통해 촬영 각도와 조명 조건을 점검하고, 앱의 사용자 수정 기능으로 오인식 결과를 보완할 예정이다."


def set_run_font(run, size: float = 10) -> None:
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Malgun Gothic")
    rfonts.set(qn("w:hAnsi"), "Malgun Gothic")
    rfonts.set(qn("w:eastAsia"), "맑은 고딕")


def replace_paragraph_text(paragraph, text: str) -> None:
    paragraph._p.clear_content()
    run = paragraph.add_run(text)
    set_run_font(run)


def main() -> None:
    doc = Document(str(INPUT))
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == OLD_TEXT:
            replace_paragraph_text(paragraph, NEW_TEXT)
            break
    else:
        raise ValueError("target paragraph not found")

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
