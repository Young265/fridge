# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
INPUT = ROOT / "output" / "hanium_dreamup_mid_report_senior_style_v6.docx"
OUTPUT = ROOT / "output" / "hanium_dreamup_mid_report_senior_style_v7.docx"


PROBLEM_ANCHOR = "- 식재료가 냉장고 내부의 정중앙에 놓인다는 보장이 없어"
SOLUTION_ANCHOR = "- YOLO 탐지 박스, OpenCV 윤곽선 후보, 중앙 fallback을 조합하여"

NEW_PROBLEM = (
    "- 현재 학습 데이터와 실제 냉장고 촬영 이미지가 충분하지 않아 일부 식재료의 "
    "인식 정확도가 낮고, 유사한 재료를 구분하는 데 한계가 있었다."
)
NEW_SOLUTION = (
    "- 식재료별 이미지 데이터를 추가 수집하고 실제 냉장고 환경에서 촬영한 데이터를 반영하여 "
    "모델 재학습과 인식 기준 보완을 진행할 예정이다."
)


def set_run_font(run, size: float = 10) -> None:
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Malgun Gothic")
    rfonts.set(qn("w:hAnsi"), "Malgun Gothic")
    rfonts.set(qn("w:eastAsia"), "맑은 고딕")


def write_paragraph(paragraph, text: str) -> None:
    paragraph._p.clear_content()
    run = paragraph.add_run(text)
    set_run_font(run)


def insert_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph.__class__(new_p, paragraph._parent)
    new_para.style = paragraph.style
    new_para.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
    new_para.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent
    new_para.paragraph_format.space_before = paragraph.paragraph_format.space_before
    new_para.paragraph_format.space_after = paragraph.paragraph_format.space_after
    new_para.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
    write_paragraph(new_para, text)
    return new_para


def find_paragraph(doc: Document, starts_with: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(starts_with):
            return paragraph
    raise ValueError(f"paragraph not found: {starts_with}")


def main() -> None:
    doc = Document(str(INPUT))

    problem_anchor = find_paragraph(doc, PROBLEM_ANCHOR)
    solution_anchor = find_paragraph(doc, SOLUTION_ANCHOR)

    if not any(p.text.strip() == NEW_PROBLEM for p in doc.paragraphs):
        insert_after(problem_anchor, NEW_PROBLEM)
    if not any(p.text.strip() == NEW_SOLUTION for p in doc.paragraphs):
        insert_after(solution_anchor, NEW_SOLUTION)

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
