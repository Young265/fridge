# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
INPUT = ROOT / "output" / "hanium_dreamup_mid_report_senior_style_v5.docx"
OUTPUT = ROOT / "output" / "hanium_dreamup_mid_report_senior_style_v6.docx"


HW_FEATURES = [
    (
        "냉장고 내부 촬영\n(Raspberry Pi 5 + 카메라)",
        "냉장고 문이 열렸을 때 내부 식재료 이미지를 촬영하고, 촬영 이미지를 AI 인식 흐름으로 전달한다.",
        "실물 장착 사진 / 첨부 예정",
    ),
    (
        "문 열림/닫힘 감지\n(리드스위치)",
        "냉장고 문 상태를 감지하여 식재료 추가 또는 소비 처리 이벤트가 실행되는 시점을 판단한다.",
        "센서 부착 사진 / 첨부 예정",
    ),
    (
        "촬영 결과 전송\n(Raspberry Pi 5 + 백엔드 연동)",
        "촬영 및 인식 결과를 서버 API로 전송하여 앱의 재고 목록에 반영될 수 있도록 한다.",
        "연동 테스트 화면 / 첨부 예정",
    ),
]


def set_run_font(run, size: float = 9.5, bold: bool = False) -> None:
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Malgun Gothic")
    rfonts.set(qn("w:hAnsi"), "Malgun Gothic")
    rfonts.set(qn("w:eastAsia"), "맑은 고딕")


def write_cell(cell, text: str, bold: bool = False, center: bool = False) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    lines = text.split("\n")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT

    for idx, line in enumerate(lines):
        if idx:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, bold=bold)


def ensure_row_count(table, target_rows: int) -> None:
    while len(table.rows) < target_rows:
        table.add_row()


def main() -> None:
    doc = Document(str(INPUT))
    table = doc.tables[11]

    ensure_row_count(table, len(HW_FEATURES) + 1)

    write_cell(table.cell(0, 0), "기능/부품", bold=True, center=True)
    write_cell(table.cell(0, 1), "설명", bold=True, center=True)
    write_cell(table.cell(0, 2), "프로젝트 실물사진", bold=True, center=True)

    for row_idx, row_values in enumerate(HW_FEATURES, start=1):
        write_cell(table.cell(row_idx, 0), row_values[0], bold=True, center=True)
        write_cell(table.cell(row_idx, 1), row_values[1])
        write_cell(table.cell(row_idx, 2), row_values[2], center=True)

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
