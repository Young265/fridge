# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
INPUT = ROOT / "output" / "hanium_dreamup_mid_report_senior_style_v3.docx"
OUTPUT = ROOT / "output" / "hanium_dreamup_mid_report_senior_style_v4.docx"


def set_run_font(run, size: float = 10, bold: bool = False) -> None:
    run.font.name = "Malgun Gothic"
    run.font.size = None
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Malgun Gothic")
    rfonts.set(qn("w:hAnsi"), "Malgun Gothic")
    rfonts.set(qn("w:eastAsia"), "맑은 고딕")


def write_paragraph(paragraph, text: str) -> None:
    paragraph._p.clear_content()
    run = paragraph.add_run(text)
    set_run_font(run)


def insert_after(paragraph, text: str, style_name: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph.__class__(new_p, paragraph._parent)
    new_para.style = style_name
    write_paragraph(new_para, text)
    return new_para


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def table_first_row_text(table) -> str:
    if not table.rows:
        return ""
    return " | ".join(cell.text.strip().replace("\n", " / ") for cell in table.rows[0].cells)


def remove_old_feature_tables(doc: Document) -> None:
    for table in list(doc.tables):
        first = table_first_row_text(table)
        if first.startswith("항목 | 기존 수동 관리 방식 | 본 프로젝트"):
            table._element.getparent().remove(table._element)
        elif first.startswith("기능 구분 | 구현 방식 | 주요 장점"):
            table._element.getparent().remove(table._element)


def update_summary_feature_cell(doc: Document) -> None:
    # The project-info summary table is part of the original template.
    info = doc.tables[4]
    cell = info.cell(9, 1)
    cell.text = ""
    p = cell.paragraphs[0]
    write_paragraph(
        p,
        "• 기존 냉장고 관리 앱처럼 사용자가 직접 입력하는 방식이 아니라, 카메라 촬영과 AI 인식으로 식재료 등록을 자동화함\n"
        "• 보유 식재료를 단순 목록으로만 보여주는 것이 아니라, 해당 재료로 만들 수 있는 레시피와 부족한 재료를 함께 안내함\n"
        "• 독거노인과 1인 가구가 식재료를 쉽게 확인하고 음식물 낭비를 줄일 수 있는 생활지원형 서비스로 확장 가능",
    )


def replace_feature_section(doc: Document) -> None:
    paragraphs = doc.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "3. 프로젝트 특·장점")
    end = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "II. 프로젝트 내용")
    style_name = paragraphs[start + 1].style.name

    for idx in range(end - 1, start, -1):
        delete_paragraph(paragraphs[idx])

    new_lines = [
        "   1) 프로젝트 주요 특징",
        "     • 본 프로젝트는 냉장고 속 식재료를 카메라와 AI로 인식하여 사용자가 직접 입력해야 하는 부담을 줄이는 스마트 냉장고 서비스이다.",
        "     • 인식된 식재료는 모바일 앱의 재고 목록으로 연결되어 사용자가 현재 보유한 재료를 쉽게 확인할 수 있다.",
        "     • 보유 재료를 기반으로 만들 수 있는 레시피와 부족한 재료를 안내하여 단순 재고 관리보다 실제 식사 준비에 더 직접적으로 도움을 준다.",
        "   2) 기존 방식과의 차별성",
        "     • 기존 냉장고 관리 앱은 사용자가 식재료명을 직접 입력하고 계속 수정해야 하지만, 본 프로젝트는 촬영과 인식 결과를 이용해 등록 과정을 자동화한다.",
        "     • 일반적인 재고 목록 서비스와 달리, 재료 확인에서 끝나지 않고 레시피 추천까지 제공하여 식재료 활용도를 높일 수 있다.",
        "     • 독거노인과 1인 가구가 냉장고 속 재료를 놓치지 않고 활용하도록 돕는 생활지원형 서비스라는 점에서 차별성이 있다.",
    ]
    anchor = paragraphs[start]
    for line in new_lines:
        anchor = insert_after(anchor, line, style_name)


def main() -> None:
    doc = Document(str(INPUT))
    remove_old_feature_tables(doc)
    update_summary_feature_cell(doc)
    replace_feature_section(doc)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
