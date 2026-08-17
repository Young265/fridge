# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from build_hanium_report import (
    ROOT,
    WORK,
    OUTPUT,
    TEMPLATE,
    add_after,
    clean_template_guidance,
    delete_row,
    ensure_output,
    insert_picture_cell,
    set_cell_text,
    shade_cell,
    write_paragraph,
)


OUT_DOCX = OUTPUT / "hanium_dreamup_mid_report_senior_style_v3.docx"
ARCH_IMG = WORK / "architecture_flow_senior_style.png"


PROJECT_NAME = "독거노인의 식생활 관리를 위한 AI 기반 스마트 냉장고"


def make_architecture_image() -> None:
    from PIL import Image, ImageDraw, ImageFont

    width, height = 1600, 410
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_path = Path(r"C:\Windows\Fonts\malgun.ttf")
    try:
        font = ImageFont.truetype(str(font_path), 30)
        small = ImageFont.truetype(str(font_path), 24)
    except Exception:
        font = ImageFont.load_default()
        small = ImageFont.load_default()

    boxes = [
        ("냉장고 내부\n카메라·리드스위치", 35, 95, 310, 265, "#E8F2FF"),
        ("Raspberry Pi\n촬영 브리지", 370, 95, 645, 265, "#EAF7EE"),
        ("YOLO 기반\n식재료 인식", 705, 95, 980, 265, "#FFF4D9"),
        ("Flask API\nMySQL 저장", 1040, 95, 1315, 265, "#F3ECFF"),
        ("Flutter 앱\n재고·레시피", 1375, 95, 1570, 265, "#FFECEC"),
    ]
    for label, x1, y1, x2, y2, fill in boxes:
        draw.rounded_rectangle([x1, y1, x2, y2], radius=24, fill=fill, outline="#344054", width=3)
        lines = label.split("\n")
        total_h = len(lines) * 39
        for idx, line in enumerate(lines):
            bbox = draw.textbbox((0, 0), line, font=font)
            tx = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
            ty = y1 + (y2 - y1 - total_h) / 2 + idx * 42
            draw.text((tx, ty), line, fill="#111827", font=font)

    for i in range(len(boxes) - 1):
        _, _x1, _y1, x2, _y2, _ = boxes[i]
        _, nx1, _ny1, _nx2, _ny2, _ = boxes[i + 1]
        y = 180
        draw.line([x2 + 18, y, nx1 - 18, y], fill="#374151", width=5)
        draw.polygon([(nx1 - 18, y), (nx1 - 40, y - 13), (nx1 - 40, y + 13)], fill="#374151")

    draw.text(
        (45, 320),
        "문 열림 감지 → 이미지 촬영 → 식재료 후보 탐지·분류 → 재고 DB 저장 → 앱에서 재고 확인 및 레시피 추천",
        fill="#4B5563",
        font=small,
    )
    ARCH_IMG.parent.mkdir(exist_ok=True)
    image.save(ARCH_IMG)


def add_matrix_after(paragraph, headings: list[str], rows: list[list[str]]) -> None:
    doc = paragraph._parent
    table = doc.add_table(rows=1, cols=len(headings), width=Inches(6.3))
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)
    for idx, heading in enumerate(headings):
        set_cell_text(table.cell(0, idx), heading, size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.cell(0, idx), "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=8.0, align=WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else None)


def add_picture_placeholder_after(paragraph):
    placeholder = WORK / "overview_picture_placeholder.png"
    if not placeholder.exists():
        from PIL import Image, ImageDraw, ImageFont

        image = Image.new("RGB", (1300, 260), "white")
        draw = ImageDraw.Draw(image)
        font_path = Path(r"C:\Windows\Fonts\malgun.ttf")
        try:
            font = ImageFont.truetype(str(font_path), 32)
            small = ImageFont.truetype(str(font_path), 22)
        except Exception:
            font = ImageFont.load_default()
            small = ImageFont.load_default()
        draw.rounded_rectangle([20, 20, 1280, 240], radius=24, outline="#94A3B8", width=4, fill="#F8FAFC")
        title = "그림 1 AI 기반 스마트 냉장고 서비스 개요도"
        sub = "냉장고 내부 촬영 화면, 앱 화면, 또는 전체 서비스 흐름도 삽입 예정"
        for idx, (line, used_font) in enumerate([(title, font), (sub, small)]):
            bbox = draw.textbbox((0, 0), line, font=used_font)
            x = (1300 - (bbox[2] - bbox[0])) / 2
            y = 82 + idx * 54
            draw.text((x, y), line, fill="#334155", font=used_font)
        image.save(placeholder)

    pic_p = add_after(paragraph, "", style="소제목")
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(str(placeholder), width=Inches(5.8))
    caption_p = add_after(pic_p, "     • 위 영역에 냉장고 내부 촬영 사진, 앱 화면, 또는 전체 서비스 흐름도를 삽입할 예정", style="소제목", size=9)
    return caption_p


def fill_summary_tables(doc: Document) -> None:
    tables = doc.tables
    set_cell_text(tables[2].cell(0, 1), PROJECT_NAME, size=11, bold=True)
    set_cell_text(tables[3].cell(0, 0), "요 약 본", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    info = tables[4]
    set_cell_text(info.cell(1, 1), PROJECT_NAME, size=9.5, bold=True)
    checks = {
        (2, 1): "☑ 생활", (2, 2): "□ 업무", (2, 3): "□ 공공/교통", (2, 4): "□ 금융/핀테크",
        (3, 1): "□ 의료", (3, 2): "□ 교육", (3, 3): "□ 유통/쇼핑", (3, 4): "□ 엔터테인먼트",
        (4, 1): "☑ 소프트웨어", (4, 2): "☑ 인공지능", (4, 3): "☑ 스마트 디바이스", (4, 4): "□ 방송, 콘텐츠",
        (5, 1): "□ 디지털융합", (5, 2): "□ 차세대통신", (5, 3): "□ 사이버보안",
    }
    for (r, c), text in checks.items():
        set_cell_text(info.cell(r, c), text, size=9)

    set_cell_text(
        info.cell(6, 1),
        "☑ 논문게재 및 포스터 발표  □ 앱등록  □ 프로그램등록  □ 특허  □ 기술이전\n"
        "□ 실용화  ☑ 공모전(한이음 드림업 공모전)  □ 기타( )",
        size=8.7,
    )
    set_cell_text(
        info.cell(7, 1),
        "• 냉장고 내부 카메라와 AI 이미지 인식 기술을 이용하여 식재료를 자동으로 등록하고 관리하는 스마트 냉장고 시스템\n"
        "• 모바일 앱을 통해 보유 식재료, 수량, 인식 이미지, 추천 레시피를 확인할 수 있도록 구현",
        size=8.6,
    )
    set_cell_text(
        info.cell(8, 1),
        "• 독거노인은 냉장고 속 식재료를 기억하기 어렵고, 식재료 방치와 음식물 폐기 문제가 발생할 수 있음\n"
        "• 자동 촬영·인식·재고화 기능을 통해 식재료 관리 부담을 낮추고 식생활 자율성을 높이고자 함",
        size=8.6,
    )
    set_cell_text(
        info.cell(9, 1),
        "• 문 열림 감지 기반 촬영, YOLO/분류 모델 기반 식재료 인식, 재고 자동 등록 및 소비 처리 기능을 통합\n"
        "• 단순 수동 입력 앱이 아니라 H/W 센서, AI 모델, 백엔드, 모바일 앱을 연결한 end-to-end 서비스 구조\n"
        "• YOLO 탐지, 윤곽선 후보, 중앙 fallback을 조합하여 식재료 위치 변화에 대응\n"
        "• 인식 결과를 재고 관리와 레시피 추천으로 연결하여 실제 식사 준비에 활용 가능",
        size=8.6,
    )
    set_cell_text(
        info.cell(10, 1),
        "• 회원가입/로그인, 냉장고 선택, 재고 CRUD, 식재료 이미지 업로드, 보유 재료 기반 레시피 추천\n"
        "• Raspberry Pi 카메라 브리지, 리드스위치 이벤트, 미리보기 스트림, 자동 업로드/소비 처리",
        size=8.6,
    )
    set_cell_text(
        info.cell(11, 1),
        "• 독거노인 및 1인 가구의 식재료 관리 편의성 향상, 음식물 폐기 감소, 스마트홈 돌봄 서비스 확장\n"
        "• AI 기반 주방 재고 관리 연구, 공모전 결과물, 복지기관 생활관리 서비스로 활용 가능",
        size=8.6,
    )
    set_cell_text(tables[5].cell(0, 0), "본    문", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def fill_body(doc: Document) -> None:
    p = doc.paragraphs
    write_paragraph(p[15], "I. 프로젝트 개요", size=13, bold=True)

    write_paragraph(p[18], "   1) 프로젝트 내용", size=10, bold=True)
    write_paragraph(
        p[19],
        "     • 본 프로젝트는 냉장고 내부 카메라와 AI 이미지 인식 기술을 활용하여 식재료를 자동으로 등록하고, 모바일 앱에서 재고와 추천 레시피를 확인할 수 있도록 하는 스마트 냉장고 시스템이다.",
        size=10,
    )
    last = p[19]
    for text in [
        "     • Raspberry Pi가 촬영·분류·업로드를 담당하고, Flask/MySQL 백엔드와 Flutter 앱이 인식 결과를 재고 관리와 레시피 추천으로 연결한다.",
    ]:
        last = add_after(last, text, style="소제목")
    last = add_picture_placeholder_after(last)

    last = add_after(last, "   2) 기획의도", style="소제목", size=10, bold=True)
    for text in [
        "     • 독거노인이 식재료를 쉽게 확인하고 식사 준비에 활용할 수 있도록, 자동 인식 기반 생활지원 흐름을 구현하고자 함.",
        "     • 단순 메모 앱이 아니라 실제 카메라 입력, AI 분석, 재고 저장, 앱 서비스를 연결한 실사용 가능한 시스템을 목표로 함.",
    ]:
        last = add_after(last, text, style="소제목")

    write_paragraph(p[23], "   1) 개발 배경 및 필요성", size=10, bold=True)
    write_paragraph(
        p[24],
        "     • 고령화와 1인 가구 증가로 독거노인의 식재료 관리와 식사 준비를 돕는 생활지원 서비스의 필요성이 커지고 있다.",
        size=10,
    )
    last = p[24]
    for text in [
        "     • 기존 수동 냉장고 관리 방식은 직접 입력 부담이 크므로, 카메라 기반 자동 인식과 앱 기반 재고 확인을 결합한 방식이 필요하다.",
    ]:
        last = add_after(last, text, style="소제목")

    last = add_after(last, "   2) 프로젝트 제작 동기 및 목적", style="소제목", size=10, bold=True)
    for text in [
        "     • AI와 IoT 기술을 독거노인의 일상 문제 해결에 적용하여 공익적 가치가 있는 생활지원 서비스를 구현하고자 함.",
        "     • 냉장고 문을 여는 자연스러운 사용 행위만으로 식재료 등록·관리·추천이 이어지는 플랫폼 구축이 목적임.",
    ]:
        last = add_after(last, text, style="소제목")

    write_paragraph(p[28], "   1) 주요 기능 중심 장점", size=10, bold=True)
    write_paragraph(
        p[29],
        "     • 본 프로젝트는 AI 이미지 인식, H/W 이벤트 감지, 재고 데이터베이스, 모바일 앱을 통합하여 실제 식재료 관리 흐름에 맞춘 서비스를 제공한다.",
        size=10,
    )
    last = p[29]
    for text in [
        "     • 일반적인 냉장고 관리 앱은 사용자가 직접 식재료를 입력하고 수량을 수정해야 하지만, 본 프로젝트는 카메라 촬영과 AI 인식을 통해 식재료 등록 과정을 자동화한다.",
        "     • 냉장고 문 열림이라는 자연스러운 사용 행동을 리드스위치로 감지하여 식재료 추가·소비 이벤트와 연결하므로, 고령층 사용자가 별도의 복잡한 조작을 하지 않아도 된다.",
        "     • 인식된 식재료는 단순히 목록에 저장되는 것이 아니라, 보유 재료 기반 레시피 추천과 부족 재료 안내로 이어져 실제 식사 준비에 바로 활용될 수 있다.",
    ]:
        last = add_after(last, text, style="소제목")

    last = add_after(last, "   2) 기존 서비스와의 차별성", style="소제목", size=10, bold=True)
    for text in [
        "     • 기존 수동 재고 관리 방식은 사용자의 기억과 입력 습관에 의존하기 때문에 장기간 유지가 어렵지만, 본 프로젝트는 촬영·인식·저장을 자동 흐름으로 구성하여 지속 사용 가능성을 높였다.",
        "     • 단순 이미지 분류만 수행하는 프로젝트와 달리, 인식 결과를 Flask API, MySQL DB, Flutter 앱 화면까지 연결하여 서비스 형태의 완성도를 확보하였다.",
        "     • 고정된 중앙 crop 방식의 한계를 줄이기 위해 YOLO 탐지 박스, OpenCV 윤곽선 후보, center fallback을 함께 사용하여 냉장고 내부 배치 변화에 대응하도록 설계하였다.",
    ]:
        last = add_after(last, text, style="소제목")

    last = add_after(last, "   3) 사용자 및 서비스 관점 장점", style="소제목", size=10, bold=True)
    for text in [
        "     • 독거노인과 1인 가구는 냉장고 속 식재료를 한눈에 확인하기 어렵기 때문에, 앱에서 재고와 추천 레시피를 함께 제공하면 식재료 활용도를 높이고 음식물 폐기를 줄일 수 있다.",
        "     • 사용자가 AI 인식 결과를 앱에서 직접 수정할 수 있도록 하여 자동화 시스템의 오류 가능성을 보완하고, 실제 사용 환경에서 신뢰성을 높일 수 있다.",
        "     • 향후 유통기한 알림, 보호자 공유, 복지기관 생활관리 플랫폼 연계 등으로 확장할 수 있어 단순 공모전 결과물을 넘어 생활지원 서비스로 발전 가능성이 있다.",
    ]:
        last = add_after(last, text, style="소제목")

    write_paragraph(p[31], "II. 프로젝트 내용", size=13, bold=True)
    write_paragraph(p[34], "   1) 서비스흐름도", size=10, bold=True)
    write_paragraph(
        p[35],
        "     • 냉장고 문 열림 감지부터 앱 표시까지 이어지는 전체 흐름은 다음과 같다.",
        size=10,
    )

    write_paragraph(p[48], "     • S/W는 Flutter 앱, Flask API, MySQL DB, YOLO 기반 인식 모듈로 구성된다.", size=10)
    write_paragraph(p[53], "     • H/W는 Raspberry Pi 5, 카메라 모듈, 리드스위치를 중심으로 구성되며 냉장고 내부 촬영 및 문 열림 이벤트 감지를 담당한다.", size=10)

    write_paragraph(p[55], "3. 주요 적용 기술", size=11, bold=True)
    write_paragraph(p[56], "     • 본 프로젝트 구현을 위해 적용한 주요 기술은 다음과 같다.", size=10)

    write_paragraph(p[60], " 5. 기타 사항 [본문에서 표현되지 못한 프로젝트의 가치(Value) 및 제작 노력]", size=10, bold=True)
    write_paragraph(p[61], "     • 본 프로젝트는 AI 인식 결과가 실제 앱 재고 관리로 연결되도록 전체 파이프라인을 직접 구현하는 데 초점을 두었다.", size=10)

    write_paragraph(p[63], "III. 프로젝트 수행 내용", size=13, bold=True)
    write_paragraph(p[65], " 1. 프로젝트 수행일정", size=11, bold=True)

    write_paragraph(p[70], "  1) 프로젝트 관리 측면", size=10, bold=True)
    write_paragraph(p[71], "     • 문제점", size=10, bold=True)
    last = p[71]
    for text in [
        "       - Flutter 앱, Flask 백엔드, AI 모델, Raspberry Pi 실행 환경이 서로 달라 실행 절차와 네트워크 주소 관리가 복잡하였다.",
        "       - H/W가 필요한 기능은 실제 장비 연결 전까지 앱과 서버만으로 동작을 확인해야 하므로 테스트 단계가 분리되는 문제가 있었다.",
        "     • 해결방안",
        "       - GitHub 저장소와 HANDOFF.md, RASPBERRY_PI.md 문서를 작성하여 백엔드 실행, Pi 연결, 테스트 명령을 표준화하였다.",
        "       - dry-run, once 실행, preview stream 기능을 마련하여 H/W 장착 전에도 인식·업로드 흐름을 단계별로 검증할 수 있도록 하였다.",
    ]:
        last = add_after(last, text, style="소제목")

    write_paragraph(p[75], "  2) 프로젝트 개발 측면", size=10, bold=True)
    write_paragraph(p[76], "     • 문제점", size=10, bold=True)
    last = p[76]
    for text in [
        "       - 식재료가 냉장고 내부의 정중앙에 놓인다는 보장이 없어 단순 중앙 crop 방식만으로는 실제 인식률이 낮아질 수 있었다.",
        "       - 문을 열 때마다 동일 식재료가 반복 등록될 수 있어 재고 데이터가 실제보다 많이 누적되는 문제가 예상되었다.",
        "     • 해결방안",
        "       - YOLO 탐지 박스, OpenCV 윤곽선 후보, 중앙 fallback을 조합하여 식재료 위치 변화에 대응하도록 후보 추출 방식을 개선하였다.",
        "       - stable frame 판정, upload cooldown, add-on-open/consume-on-close workflow를 적용하여 반복 업로드와 소비 처리를 제어하였다.",
        "       - 향후 실제 냉장고 장착 테스트와 다양한 식재료 데이터셋 추가 학습을 통해 인식 정확도를 보완할 예정이다.",
    ]:
        last = add_after(last, text, style="소제목")

    write_paragraph(p[80], "  3. 프로젝트를 통해 배우거나 느낀 점", size=11, bold=True)
    write_paragraph(
        p[81],
        "     • 이번 프로젝트를 통해 AI 모델 자체의 정확도뿐 아니라 촬영 조건, 후보 영역 추출, 서버 저장, 앱 표시까지 이어지는 전체 파이프라인 안정성이 중요함을 배웠다.",
        size=10,
    )
    last = p[81]
    for text in [
        "     • Raspberry Pi와 서버가 같은 네트워크에서 통신해야 하는 점, 카메라 FPS와 AI 추론 주기의 균형 등 H/W 연동에서 고려해야 할 요소를 경험하였다.",
        "     • 사용자 관점에서는 자동화 흐름의 신뢰성, 결과 확인의 쉬움, 오류 발생 시 수정 가능한 UI가 기능 수만큼 중요하다는 점을 확인하였다.",
    ]:
        last = add_after(last, text, style="소제목")

    write_paragraph(p[98], "Ⅳ. 기대효과 및 활용분야", size=13, bold=True)
    write_paragraph(p[99], " 1. 프로젝트의 기대효과", size=11, bold=True)
    write_paragraph(p[100], "     1. 기술적·서비스적 차별성", size=10, bold=True)
    last = p[100]
    for text in [
        "     • AI 이미지 인식 + IoT 카메라 + 모바일 앱을 결합하여 단순 수동 재고 관리보다 자동화 수준이 높다.",
        "     • 문 열림 이벤트를 활용해 사용자의 별도 입력을 줄이고, 실제 냉장고 사용 흐름에 맞는 서비스를 제공할 수 있다.",
        "     2. 사회적 기대효과",
        "     • 독거노인의 식재료 확인 부담을 줄여 식생활 자율성을 높이고, 음식물 폐기 감소에도 기여할 수 있다.",
        "     • 스마트홈, 돌봄, 생활관리 서비스와 연계하여 공공성과 확장성을 갖춘 복지형 서비스로 발전할 수 있다.",
    ]:
        last = add_after(last, text, style="소제목")

    write_paragraph(p[106], " 2. 프로젝트의 활용분야", size=11, bold=True)
    write_paragraph(p[107], "     ① 독거노인 및 1인 가구 대상 식재료 관리 서비스", size=10, bold=True)
    last = p[107]
    for text in [
        "     • 냉장고 속 식재료를 자동으로 인식하여 보유 재료와 수량을 쉽게 확인할 수 있도록 지원.",
        "     ② 지역 복지기관 및 돌봄 서비스 연계",
        "     • 고령층 생활 관리 플랫폼과 연계하여 식재료 보유 현황, 식사 준비 지원, 돌봄 알림 서비스로 확장 가능.",
        "     ③ AI 기반 스마트 주방 연구 및 교육",
        "     • 이미지 인식, IoT 센서, 모바일 앱, 백엔드 API를 결합한 융합 프로젝트 사례로 활용 가능.",
    ]:
        last = add_after(last, text, style="소제목")


def fill_project_tables(doc: Document) -> None:
    tables = doc.tables
    make_architecture_image()

    arch_cell = tables[11].cell(0, 0)
    arch_cell.text = ""
    title = arch_cell.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_paragraph(title, "그림 1 AI 기반 스마트 냉장고 서비스 시스템 구성도", size=9.5, bold=True)
    pic = arch_cell.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(ARCH_IMG), width=Inches(6.2))

    functions = tables[12]
    rows = [
        ("S/W", "식재료 이미지 인식", "카메라 이미지에서 식재료 후보를 탐지하고 YOLO/분류 모델로 식재료명을 판별하는 기능\n8/31까지 클래스 확장 및 정확도 검증 예정", "75%"),
        ("S/W", "재고 관리 앱", "냉장고별 식재료 목록 조회, 수량·단위·상태 수정, 사용자 직접 추가/삭제 기능\n8/20까지 UI 보완 예정", "70%"),
        ("S/W", "레시피 추천", "보유 식재료와 레시피 필요 재료를 비교하여 만들 수 있는 레시피와 부족 재료를 제공하는 기능\n9/10까지 추천 데이터 보강 예정", "60%"),
        ("H/W", "카메라·문 열림 감지", "Raspberry Pi 카메라 촬영, 리드스위치 기반 문 열림/닫힘 이벤트 감지 및 업로드/소비 workflow\n9/10까지 실물 장착 테스트 예정", "55%"),
    ]
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            set_cell_text(functions.cell(r, c), value, size=8.2 if c == 2 else 8.8)
    delete_row(functions, 5)

    sw_table = tables[13]
    sw_rows = [
        ("식재료 이미지 인식", "냉장고 내부 카메라로 촬영한 이미지를 분석하여 식재료 후보 영역을 추출하고 라벨과 신뢰도를 산출"),
        ("재고 관리", "인식된 식재료를 냉장고별 재고로 저장하고 앱에서 수량, 단위, 상태, 메모를 수정 가능"),
        ("레시피 추천", "현재 보유 식재료를 기준으로 만들 수 있는 요리와 부족한 식재료를 함께 안내"),
    ]
    images = [
        ROOT / "fridge_app" / "assets" / "images" / "sample_fruit_01.jpg",
        ROOT / "backend" / "egg_dataset_preview.jpg",
        ROOT / "fridge_app" / "assets" / "images" / "sample_fruit_02.jpg",
    ]
    captions = ["식재료 촬영 예시", "AI 학습·검증 이미지", "앱 표시용 샘플 이미지"]
    for idx, (name, desc) in enumerate(sw_rows, start=1):
        set_cell_text(sw_table.cell(idx, 0), name, size=8.4, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(sw_table.cell(idx, 1), desc, size=8.1)
        insert_picture_cell(sw_table.cell(idx, 2), images[idx - 1], captions[idx - 1], width=1.1)
    delete_row(sw_table, 4)

    hw_table = tables[15]
    hw_rows = [
        ("Raspberry Pi 5 + 카메라", "냉장고 내부 식재료 촬영, AI 인식 파이프라인 실행, 백엔드 업로드를 담당"),
        ("리드스위치", "냉장고 문 열림/닫힘을 감지하여 식재료 추가 및 소비 이벤트를 자동 트리거"),
    ]
    for idx, (name, desc) in enumerate(hw_rows, start=1):
        set_cell_text(hw_table.cell(idx, 0), name, size=8.4, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(hw_table.cell(idx, 1), desc, size=8.1)
        set_cell_text(hw_table.cell(idx, 2), "실물 장착 사진\n첨부 예정", size=8.4, align=WD_ALIGN_PARAGRAPH.CENTER)

    tech = (
        "1. 딥러닝 기반 식재료 이미지 인식\n"
        "적용 기술: Ultralytics YOLOv8n, 식재료 분류 모델(best.pt)\n"
        "적용 목적: 냉장고 내부 이미지에서 식재료 후보를 자동으로 탐지하고 식재료명을 분류\n"
        "이론 기반: 객체 탐지 모델은 이미지 내 후보 영역을 bounding box로 예측하고, 분류 모델은 crop 이미지의 class probability를 산출\n"
        "적용 방식: YOLO 탐지 박스 → crop padding → 분류 모델 추론 → label/confidence 산출 → Flask /upload 전송\n\n"
        "2. OpenCV 기반 후보 영역 보완\n"
        "적용 기술: Canny edge, contour detection, center fallback\n"
        "적용 목적: 일반 탐지 모델이 알지 못하는 식재료나 배치가 불규칙한 상황에서도 후보 영역을 확보\n"
        "적용 방식: detector 후보가 부족하면 윤곽선 기반 후보를 추가하고, 마지막으로 중앙 crop을 fallback으로 사용\n\n"
        "3. Flask REST API 및 MySQL 재고 관리\n"
        "적용 기술: Flask, PyMySQL, MySQL, CORS\n"
        "적용 목적: 사용자, 냉장고, 식재료, 재고, 레시피, 이미지 업로드 데이터를 서버에서 통합 관리\n"
        "적용 방식: /upload, /consume, /inventory, /recipes API를 통해 자동 인식 결과와 앱 화면을 연결\n\n"
        "4. Flutter 기반 모바일 앱\n"
        "적용 기술: Flutter, Dart, HTTP 통신, Material UI\n"
        "적용 목적: 사용자가 냉장고별 재고와 추천 레시피를 쉽게 확인하고 직접 수정할 수 있는 인터페이스 제공\n"
        "적용 방식: 로그인/회원가입 → 냉장고 선택 → 재고 목록 → 상세 수정 → 레시피 추천 순서로 화면 구성\n\n"
        "5. Raspberry Pi 기반 IoT 연동\n"
        "적용 기술: Picamera2/OpenCV camera, gpiozero, reed switch, MJPEG preview stream\n"
        "적용 목적: 냉장고 문 열림 이벤트와 카메라 촬영을 연결하여 식재료 추가·소비 이벤트를 자동화\n"
        "적용 방식: 문 열림 시 add scan, 문 닫힘 시 consume scan을 실행하고 stable frame 판정 후 백엔드로 전송"
    )
    set_cell_text(tables[18].cell(0, 0), tech, size=8.9)

    env = tables[19]
    env_values = {
        1: "Windows 개발 PC, Raspberry Pi OS 64-bit, Android/Web 테스트 환경",
        2: "Visual Studio Code, Flutter 개발 환경, Python 가상환경",
        3: "Flutter, Flask, MySQL, Ultralytics YOLO, OpenCV, Git/GitHub",
        4: "Dart, Python, SQL",
        5: "REST API, CORS, YOLO 모델 가중치, 이미지 업로드 저장소",
        6: "Raspberry Pi 5, 카메라 모듈, 냉장고, Android 스마트폰",
        7: "리드스위치(GPIO17), 카메라",
        8: "HTTP REST 통신, LAN/Tailscale 기반 연결, MJPEG preview stream",
        9: "Python",
        10: "640x480 30FPS 기본 촬영, warm/on-demand camera 모드 지원",
        11: "GitHub, HANDOFF.md, RASPBERRY_PI.md",
        12: "카카오톡, 디스코드, 대면 미팅, Notion 활용 예정",
        13: "주 1회 진행 공유, 기능 단위 테스트 및 백업",
    }
    for row_idx, value in env_values.items():
        set_cell_text(env.cell(row_idx, 2), value, size=8.2)
    delete_row(env, 14)

    value = (
        "• 프로젝트를 통한 가치창출\n"
        "  - AI 기술과 IoT 센서를 결합하여 독거노인의 식생활 관리라는 생활 문제를 해결하는 사회적 가치 실현\n"
        "  - 식재료 인식, 재고 관리, 레시피 추천 결과를 하나의 정보 구조로 제공하여 사용자의 식사 준비 시간을 단축\n\n"
        "• 차별화된 성능\n"
        "  - 고정 중앙 crop이 아닌 YOLO 탐지, 윤곽선 후보, center fallback을 조합하여 실제 냉장고 배치 변화에 대응\n"
        "  - stable frame, cooldown, add/consume workflow를 적용하여 반복 업로드와 재고 오차를 줄임\n\n"
        "• 신뢰성\n"
        "  - Flask API와 MySQL DB를 통해 인식 결과, 원본 이미지, crop 이미지, confidence를 함께 저장\n"
        "  - dry-run, once 실행, preview stream으로 H/W 장착 전후를 나누어 단계별 검증 가능\n\n"
        "• 사용성\n"
        "  - 사용자가 직접 입력하지 않아도 식재료가 자동 등록되며, 앱에서 결과를 확인하고 수정할 수 있음\n"
        "  - 냉장고별 재고와 레시피 추천을 제공하여 독거노인과 1인 가구의 식재료 활용도를 높임\n\n"
        "• 데이터와 프로그램의 가치\n"
        "  - 식재료 이미지, 인식 confidence, 재고 변화 데이터를 축적하여 향후 모델 개선과 식생활 패턴 분석에 활용 가능"
    )
    set_cell_text(tables[20].cell(0, 0), value, size=8.9)

    schedule = tables[22]
    set_cell_text(schedule.cell(1, 0), "작성 기준", size=8.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(schedule.cell(1, 1), "2026. 7. 7. 기준", size=8.2)
    for c in range(2, 11):
        set_cell_text(schedule.cell(1, c), "", size=8)

    plans = {
        4: ("계획", "프로젝트 주제 선정 및 목표 설정", [2, 3]),
        5: ("분석", "독거노인 식재료 관리 문제 및 요구사항 분석", [2, 3, 4]),
        6: ("설계", "시스템 구조, DB/API 설계", [3, 4]),
        7: ("설계", "앱 UI 및 Raspberry Pi 연동 구조 설계", [3, 4, 5]),
        8: ("개발", "Flask/MySQL 백엔드 및 재고 API 구현", [4, 5]),
        9: ("개발", "Flutter 앱 화면 및 API 연동", [4, 5, 6]),
        10: ("개발", "AI 식재료 인식 및 카메라 브리지 연동", [5, 6, 7]),
        11: ("테스트", "기능 테스트, 인식 정확도 검증, H/W 장착 테스트", [6, 7, 8]),
        12: ("종료\n(성과등록)", "최종 보고서 작성 및 발표 자료 정리", [8]),
    }
    for r, (category, task, months) in plans.items():
        set_cell_text(schedule.cell(r, 0), category, size=8.1, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(schedule.cell(r, 1), task, size=7.9)
        for c in range(2, 11):
            mark = "완료" if c in months and c <= 4 else ("진행" if c in months and c == 5 else ("예정" if c in months else ""))
            set_cell_text(schedule.cell(r, c), mark, size=7.2, align=WD_ALIGN_PARAGRAPH.CENTER)
            if mark == "완료":
                shade_cell(schedule.cell(r, c), "E2F0D9")
            elif mark == "진행":
                shade_cell(schedule.cell(r, c), "FFF2CC")
            elif mark == "예정":
                shade_cell(schedule.cell(r, c), "EAF2F8")

    set_cell_text(
        tables[28].cell(0, 0),
        "효과 영역별 기대효과\n"
        "정보 접근성 향상: 냉장고 내부 식재료 현황을 앱에서 쉽게 확인할 수 있어 기억 의존도를 낮춤\n"
        "사용자 자율성 강화: 직접 입력 없이 식재료가 등록되어 고령층도 부담 없이 재고 관리 가능\n"
        "사회적 포용 실현: 독거노인과 1인 가구의 식생활 관리를 지원하는 생활밀착형 AI 서비스로 활용 가능\n"
        "서비스 확장 기반 마련: 향후 유통기한 알림, 식단 추천, 보호자/복지기관 연계 서비스로 확장 가능",
        size=8.9,
    )
    set_cell_text(
        tables[29].cell(0, 0),
        "실질적인 활용분야\n"
        "① 독거노인 및 1인 가구 식재료 관리 앱\n"
        "② 지역 복지기관·돌봄 서비스의 생활 관리 보조 시스템\n"
        "③ AI 이미지 인식 및 IoT 융합 교육·연구 사례\n"
        "④ 스마트 냉장고·스마트 주방 플랫폼의 기능 모듈",
        size=8.9,
    )


def add_strength_tables(doc: Document) -> None:
    anchor = doc.paragraphs[29]
    add_matrix_after(
        anchor,
        ["기능 구분", "구현 방식", "주요 장점"],
        [
            ["자동 인식", "카메라 촬영 + YOLO/분류 모델", "수동 입력 없이 식재료 등록 가능"],
            ["문 열림 감지", "리드스위치 GPIO 이벤트", "냉장고 사용 흐름과 자연스럽게 연동"],
            ["재고 관리", "Flask API + MySQL + Flutter 앱", "냉장고별 재고 조회·수정 가능"],
            ["레시피 추천", "보유 식재료와 필요 재료 비교", "식재료 활용도 향상 및 폐기 감소"],
        ],
    )
    add_matrix_after(
        anchor,
        ["항목", "기존 수동 관리 방식", "본 프로젝트"],
        [
            ["입력 방식", "사용자가 직접 식재료명 입력", "카메라 촬영 후 AI 자동 인식"],
            ["사용 편의성", "지속적인 수기 관리 필요", "문 열림 이벤트 기반 자동화 가능"],
            ["서비스 범위", "단순 목록 관리 중심", "재고 관리 + 레시피 추천 + 소비 처리"],
        ],
    )


def main() -> None:
    ensure_output()
    doc = Document(str(TEMPLATE))
    fill_summary_tables(doc)
    fill_body(doc)
    fill_project_tables(doc)
    add_strength_tables(doc)
    clean_template_guidance(doc)
    doc.save(str(OUT_DOCX))
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
