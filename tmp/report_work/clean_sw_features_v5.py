# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
INPUT = ROOT / "output" / "hanium_dreamup_mid_report_senior_style_v4.docx"
OUTPUT = ROOT / "output" / "hanium_dreamup_mid_report_senior_style_v5.docx"


SW_FEATURES = [
    (
        "식재료 이미지 인식",
        "냉장고 내부 카메라로 촬영한 이미지를 분석하여 식재료명을 인식하고 인식 결과를 서버로 전달한다.",
    ),
    (
        "재고 목록 조회 및 관리",
        "인식된 식재료를 냉장고별 재고 목록으로 저장하고, 앱에서 보유 식재료와 수량을 확인할 수 있도록 한다.",
    ),
    (
        "사용자 직접 수정",
        "AI 인식 결과가 틀리거나 누락된 경우 사용자가 앱에서 식재료를 직접 추가, 수정, 삭제할 수 있도록 한다.",
    ),
    (
        "레시피 추천",
        "현재 보유한 식재료를 기준으로 만들 수 있는 레시피를 추천하고, 부족한 재료를 함께 안내한다.",
    ),
    (
        "서버 및 데이터 연동",
        "Flask API와 MySQL DB를 통해 이미지 업로드, 재고 조회, 소비 처리, 레시피 데이터를 통합 관리한다.",
    ),
]


def set_run_font(run, size: float = 9.5, bold: bool = False) -> None:
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Malgun Gothic")
    rfonts.set(qn("w:hAnsi"), "Malgun Gothic")
    rfonts.set(qn("w:eastAsia"), "맑은 고딕")


def write_cell(cell, text: str, bold: bool = False, align_center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    if align_center:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph._p.clear_content()
    run = paragraph.add_run(text)
    set_run_font(run, size=10)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "8A8A8A")


def set_table_geometry(table, widths: list[int]) -> None:
    total_width = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is not None:
        table._tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    table._tbl.insert(0, grid)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(table, margin_dxa: int = 100) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for edge in ("top", "left", "bottom", "right"):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(margin_dxa))
        element.set(qn("w:type"), "dxa")


def build_sw_table(doc: Document):
    table = doc.add_table(rows=len(SW_FEATURES) + 1, cols=2)
    table.autofit = False
    set_table_geometry(table, [2200, 6800])
    set_cell_margins(table)
    set_table_borders(table)

    write_cell(table.cell(0, 0), "기능", bold=True, align_center=True)
    write_cell(table.cell(0, 1), "설명", bold=True, align_center=True)
    set_cell_shading(table.cell(0, 0), "D9EAF7")
    set_cell_shading(table.cell(0, 1), "D9EAF7")

    for row_idx, (feature, description) in enumerate(SW_FEATURES, start=1):
        write_cell(table.cell(row_idx, 0), feature, bold=True, align_center=True)
        write_cell(table.cell(row_idx, 1), description)

    return table


def main() -> None:
    doc = Document(str(INPUT))

    old_sw_table = doc.tables[10]
    anchor = next(
        p
        for p in doc.paragraphs
        if p.text.strip().startswith("• S/W는 Flutter 앱")
    )
    set_paragraph_text(
        anchor,
        "• S/W는 Flutter 앱, Flask API, MySQL DB, AI 인식 모듈로 구성되며 주요 기능은 다음과 같다.",
    )

    new_table = build_sw_table(doc)
    anchor._p.addnext(new_table._element)
    old_sw_table._element.getparent().remove(old_sw_table._element)

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
